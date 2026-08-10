from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches


SOURCE = Path(r"D:\workspace\demo_project\multi-agent-hot-copy-generator\中文简历模板-标题图片位置已调整.docx")
PHOTO = Path(r"E:\work\郭雯白底jpg.jpg")
OUTPUT = Path(r"D:\workspace\demo_project\multi-agent-hot-copy-generator\郭雯简历-已插入证件照.docx")

doc = Document(SOURCE)
photo_paragraph = doc.paragraphs[2]

# Remove placeholder text while preserving the floating frame paragraph.
for run in list(photo_paragraph.runs):
    photo_paragraph._p.remove(run._r)

# Remove placeholder-only fill and border.
p_pr = photo_paragraph._p.get_or_add_pPr()
for tag in ("shd", "pBdr"):
    element = p_pr.find(qn(f"w:{tag}"))
    if element is not None:
        p_pr.remove(element)

# Let the image's natural 3:4 aspect ratio determine the frame height.
frame = p_pr.find(qn("w:framePr"))
if frame is not None:
    frame.set(qn("w:w"), "1300")
    frame.set(qn("w:y"), "-300")
    frame.attrib.pop(qn("w:h"), None)
    frame.attrib.pop(qn("w:hRule"), None)

run = photo_paragraph.add_run()
shape = run.add_picture(str(PHOTO), width=Inches(0.9))
shape._inline.docPr.set("descr", "郭雯白底证件照")

doc.save(OUTPUT)
print(OUTPUT)
