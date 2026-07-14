"""Convert docs/resume_agent_fullstack.md to docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "resume_agent_fullstack.md"
DOCX_PATH = ROOT / "docs" / "resume_agent_fullstack.docx"


def set_cn_font(run, name: str = "微软雅黑", size: int | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_rich_paragraph(doc: Document, text: str, style: str | None = None, size: int = 10.5):
    para = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_cn_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_cn_font(run, size=size)
    return para


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = [line.strip() for line in lines if line.strip().startswith("|")]
    if len(rows) < 2:
        return [], []
    headers = [c.strip() for c in rows[0].strip("|").split("|")]
    body = []
    for row in rows[2:]:
        body.append([c.strip() for c in row.strip("|").split("|")])
    return headers, body


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, bold=True, size=10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_cn_font(run, size=10)


def convert(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            set_cn_font(run, size=16, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            set_cn_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("### "):
            add_rich_paragraph(doc, stripped[4:].strip(), size=11)
            doc.paragraphs[-1].runs[0].bold = True
            set_cn_font(doc.paragraphs[-1].runs[0], size=11, bold=True)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            headers, rows = parse_table(block)
            if headers:
                add_table(doc, headers, rows)
                doc.add_paragraph()
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:].strip(), size=10)
            doc.paragraphs[-1].paragraph_format.left_indent = Pt(12)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            content = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*[^*]+\*\*)", content)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, size=10.5, bold=True)
                else:
                    run = p.add_run(part)
                    set_cn_font(run, size=10.5)
            i += 1
            while i < len(lines) and lines[i].startswith("   "):
                add_rich_paragraph(doc, lines[i].strip(), size=10.5)
                doc.paragraphs[-1].paragraph_format.left_indent = Pt(18)
                i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped[2:])
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, size=10.5, bold=True)
                else:
                    run = p.add_run(part)
                    set_cn_font(run, size=10.5)
            i += 1
            continue

        add_rich_paragraph(doc, stripped, size=10.5)
        i += 1

    return doc


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = convert(md_text)
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
