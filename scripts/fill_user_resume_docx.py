"""Fill user resume template with project content (bottom-up to preserve indices)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TEMPLATE = Path(r"E:\录屏截屏\备忘录\ai全栈工程师简历.docx")
OUTPUT = Path(r"E:\录屏截屏\备忘录\ai全栈工程师简历_已填写.docx")
BACKUP = TEMPLATE.with_suffix(".docx.bak")


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_para_text(paragraph, text: str, bold: bool | None = None):
    paragraph.clear()
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold


def set_title_line(paragraph, title: str, date: str, title_bold: bool = True):
    paragraph.clear()
    r1 = paragraph.add_run(title)
    r1.bold = title_bold
    paragraph.add_run("\t")
    paragraph.add_run(date)


def add_bullets(after_para, items: list[str]) -> Paragraph:
    last = after_para
    for item in items:
        last = insert_paragraph_after(last, item, style="List Bullet")
    return last


def add_numbered_items(after_para, items: list[str]) -> Paragraph:
    last = after_para
    for idx, item in enumerate(items, start=1):
        last = insert_paragraph_after(last, f"{idx}. {item}", style="List Number")
    return last


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)

    shutil.copy2(TEMPLATE, BACKUP)
    doc = Document(str(TEMPLATE))
    p = doc.paragraphs

    # ----- 从下往上改，避免 insert 打乱尚未处理的段落索引 -----

    # 教育背景
    set_para_text(p[20], "吉林工程技术师范学院 · 软件工程 · 本科")

    # 项目2 亮点
    set_para_text(p[17], "项目亮点：")
    highlights2 = [
        "完成部门项目统计、个人信息、部门信息管理等模块的前端开发与联调。",
        "封装列表、表单等通用组件，提升同类业务页开发效率。",
    ]
    add_numbered_items(p[17], highlights2)

    set_para_text(
        p[16],
        "项目描述：Vue 开发的网页端社区工作人员端与面向居民的小程序客户端；负责页面开发、接口联调与模块交付。",
    )
    set_title_line(p[15], "智慧社区客户端政府端", "2019年3月 – 2019年6月", title_bold=True)

    # 项目1 亮点
    set_para_text(p[14], "项目亮点：")
    highlights1 = [
        "三 Agent 顺序编排 + 多级降级容错：需求理解 → 文案创作 → 审核优化；单 Agent 失败不导致整单失败，状态与 Token 落库复盘。",
        "ReAct Function Calling + 11 个可插拔 Skill：BaseAgent 循环引擎 + SkillRegistry；max_tool_calls 防止工具调用死循环。",
        "LangGraph 双 StateGraph 驱动头条 RAG：ingest 图 chunk→index，query 图 retrieve→format；600 字/块、80 字重叠中文切分。",
        "双路 RAG + 热榜闭环：历史爆款向量库 + 头条长文参考库并行增强创作；APScheduler 每小时同步热榜并向量化。",
        "全栈交付：FastAPI REST + JWT 鉴权；React AgentPipeline 可视化三阶段；Agent/Tool 调用链落库可排查。",
        "编排引擎抽象 OrchestrationEngine，支持 native / LangGraph 编排切换，预留扩展点。",
    ]
    add_numbered_items(p[14], highlights1)

    desc1 = (
        "项目描述：面向营销/内容场景的 Multi-Agent 文案生产平台。用户提交需求后，系统自动匹配热榜话题、"
        "检索历史爆款与头条参考文风，经三阶段 Agent 流水线输出可发布文案；支持任务追踪、Agent 日志与热榜管理。"
        "技术栈：Python · FastAPI · DeepSeek · LangGraph · LangChain · ChromaDB · MySQL · React · JWT · Docker。"
    )
    set_para_text(p[13], desc1)
    set_title_line(p[12], "多Agent 热点爆款文案生成系统", "2024年 – 2025年", title_bold=True)

    # 工作经历
    set_para_text(p[10], "系统学习 Python Web、LLM 应用与 Agent 架构；独立完成多智能体热点文案生成系统从 0 到 1。")
    set_title_line(p[9], "自主转型学习（AI 应用方向）", "2022年6月 – 2024年6月", title_bold=False)
    set_para_text(p[8], "Vue 2/3 后台与业务页开发；接口联调、组件封装、表单列表与基础性能优化。")
    set_title_line(p[7], "XX科技-前端开发工程师", "2024年7月 – 2025年7月")

    # 专业技能
    set_para_text(
        p[5],
        "前端基础：1 年 Vue 业务开发（组件化、路由、Axios、Element UI）；独立使用 React 18 + TypeScript + Vite 完成 Agent 任务管理前端（鉴权、Pipeline 可视化）。",
    )
    skills = [
        "后端与工程：FastAPI REST API、SQLAlchemy 2.0、MySQL 建模、JWT 鉴权、APScheduler 定时任务、Docker Compose；熟悉 API → Service → Agent/Skill 分层。",
        "Prompt 与 Agent 工程：多 Agent System Prompt 设计；ReAct 式 Function Calling 循环（模型决策 → Tool 执行 → 结果回注）；理解「生成 → 执行 → 校验 → 降级」闭环。",
        "AI Agent 开发：3 Agent 顺序编排 + 11 个可插拔 Skill；LangGraph StateGraph 构建 RAG 入库/检索双图；Chroma 向量检索 + 本地 Embedding；热榜自动同步。",
        "AI 辅助开发：熟练使用 Cursor 进行需求拆解、代码生成、重构与 Debug。",
    ]
    add_bullets(p[5], skills)

    # 标题
    set_para_text(p[0], "郭雯 - AI全栈工程师", bold=True)

    out_path = OUTPUT
    try:
        doc.save(str(TEMPLATE))
        out_path = TEMPLATE
        print(f"Updated: {TEMPLATE}")
    except PermissionError:
        doc.save(str(OUTPUT))
        print(f"原文件被占用，已另存为: {OUTPUT}")
        print("请关闭 Word 后，可将新文件覆盖原文件，或直接投递新文件。")
    print(f"Backup:  {BACKUP}")


if __name__ == "__main__":
    main()
