from pathlib import Path

from docx import Document


path = Path(r"D:\workspace\demo_project\multi-agent-hot-copy-generator\中文简历模板-调整版.docx")
doc = Document(path)

print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index}: style={paragraph.style.name!r} align={paragraph.alignment!r} text={text!r}")

for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        values = [" / ".join(p.text.strip() for p in cell.paragraphs if p.text.strip()) for cell in row.cells]
        print(f"  R{row_index}: {values}")

for index, shape in enumerate(doc.inline_shapes):
    print(f"SHAPE {index}: width={shape.width} height={shape.height} type={shape.type}")
